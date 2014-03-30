import xlrd
from docx import Document


class HoroDocument():

	def generate(self, content):
		self.doc = Document()
		for item in content:
			self.doc.add_heading(item['section_title'], level=1)
			for para in item['paragraph']:
				self.doc.add_heading(para['title'], level=3)
				self.doc.add_paragraph(para['content'])

	def save(self, filename):
		self.doc.save(filename)


# Load content using input
class HoroReport():

	def __init__(self):
		self._data = [] 

	def loadDataSource(self, fileName):
		self._workbook = xlrd.open_workbook(fileName)

	def loadContentUsing(self, reportData):
		content = []
		for item in reportData:
			if item['type'] == 'Section':
				section = {
					'section_title': item['value']
				}
				section['paragraph'] = []
				content.append(section)
			else: 
				if item['type'] == 'Paragraph':
					sheetName = item['value']['sheet']
					contentRow = item['value']['value']
					paragraph = self.getParagraphFromSource(sheetName, contentRow)
					section['paragraph'].append(paragraph)
		return content

	def getParagraphFromSource(self, sheetName, value):
		notFoundResult = {'title': 'Cannot found', 'content': 'Cannot found'}
		try:
			sheet = self._workbook.sheet_by_name(str(sheetName))
			for row in range(sheet.nrows):
				if sheet.cell_value(row, 0) == value:
					return {
						'title': sheet.cell_value(row, 1),
						'content': sheet.cell_value(row, 2)
						}
		except Exception as e:
			return notFoundResult
		return notFoundResult


# Parse input from xlsx file
class HoroInput():

	def load(self, filename):
		self._workbook = xlrd.open_workbook(filename)
		self._reportNames = self._workbook.sheet_names()
		self._report = {}
		for name in self._reportNames:
			self._report[name] = self._loadSection(name)

	def _isValidParagraphInput(self, row):
		if row[0] is not None and row[1] is not None:
			return True
		else:
			return False

	def clean(self, data):
		res = []
		for item in data:
			if item[0] == 'Section':
				res.append({'Section': item[1]})
			else:
				if self._isValidParagraphInput(item):
					res.append({'Paragraph': {
							'sheet': item[0],
							'value': item[1]
						}})
		return res

	def getReportNames(self):
		return self._reportNames

	def getReport(self, name):
		return self._report[name]

	def _loadSection(self, sheet_name):
		data = []
		sheet = self._workbook.sheet_by_name(sheet_name)
		for row in range(sheet.nrows):
			firstCol = sheet.cell_value(row, 0)
			secondCol = sheet.cell_value(row, 1)
			if firstCol == 'Section':
				typeValue = 'Section'
				dataValue = secondCol
			else:
				typeValue = 'Paragraph'
				dataValue = {
					'sheet': firstCol,
					'value': secondCol
				}
			data.append({
				'type': typeValue,
				'value': dataValue
				})
		return data

class ExcelReader():

	def __init__(self, filename):
		self._workbook = xlrd.open_workbook(filename)

	def getSheets(self):
		return self._workbook.sheet_names()	

	def getSheetContent(self, sheetName):
		sheet = self._workbook.sheet_by_name(sheetName)
		content = []
		for row in range(sheet.nrows):
			rowContent = []
			for col in range(sheet.ncols):
				rowContent.append(sheet.cell_value(row, col))
			content.append(rowContent)			
		return content
