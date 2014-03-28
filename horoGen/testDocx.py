# coding=utf-8
from docx import Document
import xlrd
import unittest

class TestHoroInput(unittest.TestCase):

	def setUp(self):
		self.input = HoroInput()
		self.input.load('input.xlsx')

	def testSanity(self):
		self.assertEqual(['Sheet1'], self.input.getReportNames())
		
	def testGetParagraphInfo(self):
		self.assertEqual('Section', self.input.getType(0))
		self.assertEqual(u'Đường đi', self.input.getValue(0))
		self.assertEqual('Paragraph', self.input.getType(1))
		self.assertDictEqual({'sheet': 1.1, 'value': 1}, self.input.getValue(1))


class HoroInput():

	def __init__(self):
		self._reportNames = []

	def load(self, filename):
		self._workbook = xlrd.open_workbook(filename)
		self._reportNames = self._workbook.sheet_names()
		for name in self._reportNames:
			self._loadSection(name)

	def getInputCount(self):
		return len(self._rawInput)

	def getType(self, row):
		return self._rawInput[row]['type']

	def getValue(self, row):
		return self._rawInput[row]['value']

	def _loadSection(self, sheet_name):
		self._rawInput = []
		sheet = self._workbook.sheet_by_name(sheet_name)
		for row in range(sheet.nrows):
			firstCol = sheet.cell_value(row, 0)
			secondCol = sheet.cell_value(row, 1)
			if firstCol == 'Section':
				typeValue = 'Section'
				dataValue = secondCol
			else:
				typeValue = 'Paragraph'
				dataValue = {
					'sheet': firstCol,
					'value': secondCol
				}
			self._rawInput.append({
				'type': typeValue,
				'value': dataValue
				})

	def getReportNames(self):
		return self._reportNames
		

class ExcelReader():

	def load(self, filename):
		workbook = xlrd.open_workbook(filename)
		sheet = workbook.sheet_by_name('1.1')
		return sheet.cell_value(1,2)


class TestDocx(unittest.TestCase):

	def testSanity(self):
		doc = Document()	
		excel = ExcelReader()
		content = excel.load('AstroReport_Content.xlsx')
		heading = doc.add_heading('Heading, level 1', level=1)
		doc.add_paragraph(content)
		doc.save('demo.docx')

if __name__ == '__main__':
	unittest.main()


